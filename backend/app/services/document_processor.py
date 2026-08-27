import io
import re
import logging
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field
import pymupdf  # PyMuPDF
import docx

logger = logging.getLogger(__name__)

@dataclass
class DocumentPage:
    """Represents an individual page of extracted text."""
    page_number: int
    text: str

@dataclass
class DocumentContent:
    """Structured representation of an extracted document with preserved pages."""
    title: str
    pages: List[DocumentPage] = field(default_factory=list)
    total_pages: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)

class DocumentProcessor:
    """Service for extracting, normalizing, and cleaning text from PDF and DOCX technical documents."""

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Normalizes extracted text:
        - Replaces non-breaking and unusual spaces with standard space
        - Removes excessive consecutive blank lines and trailing whitespace
        - Strips control characters while preserving structural paragraphs
        """
        if not text:
            return ""
        
        # Normalize non-standard whitespace and unicode spaces
        text = re.sub(r"[\r\f\v]", "\n", text)
        text = re.sub(r"[\t\u00A0\u2000-\u200B\u202F\u205F\u3000]", " ", text)
        
        # Collapse 3+ newlines into 2
        text = re.sub(r"\n{3,}", "\n\n", text)
        
        # Clean lines
        lines = [line.strip() for line in text.split("\n")]
        cleaned_text = "\n".join(lines).strip()
        return cleaned_text

    def extract_from_pdf(
        self,
        file_bytes: bytes,
        title: str = "",
        metadata: Optional[Dict[str, Any]] = None
    ) -> DocumentContent:
        """
        Extracts text from PDF page by page using PyMuPDF.
        Preserves 1-indexed page numbers.
        """
        pages: List[DocumentPage] = []
        doc = None
        try:
            doc = pymupdf.open(stream=file_bytes, filetype="pdf")
            total_pages = len(doc)
            logger.info(f"Extracting text from PDF: '{title}' ({total_pages} pages)")

            for page_idx in range(total_pages):
                page = doc.load_page(page_idx)
                raw_text = page.get_text("text") or ""
                cleaned = self.clean_text(raw_text)
                
                if cleaned:
                    pages.append(DocumentPage(page_number=page_idx + 1, text=cleaned))
                else:
                    # Still record empty page placeholder to preserve page numbering continuity
                    pages.append(DocumentPage(page_number=page_idx + 1, text=""))

            return DocumentContent(
                title=title,
                pages=pages,
                total_pages=total_pages,
                metadata=metadata or {}
            )
        except Exception as e:
            logger.error(f"Failed to extract text from PDF '{title}': {e}", exc_info=True)
            raise ValueError(f"Error parsing PDF document: {str(e)}")
        finally:
            if doc:
                doc.close()

    def extract_from_docx(
        self,
        file_bytes: bytes,
        title: str = "",
        metadata: Optional[Dict[str, Any]] = None
    ) -> DocumentContent:
        """
        Extracts text from DOCX documents using python-docx.
        Simulates page boundaries based on section/paragraph density.
        """
        try:
            doc_file = io.BytesIO(file_bytes)
            doc = docx.Document(doc_file)
            logger.info(f"Extracting text from DOCX: '{title}' ({len(doc.paragraphs)} paragraphs)")

            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            # Group into simulated pages (approx 450 words or section breaks)
            pages: List[DocumentPage] = []
            current_page_text: List[str] = []
            current_word_count = 0
            page_number = 1

            for p in paragraphs:
                current_page_text.append(p)
                current_word_count += len(p.split())
                
                if current_word_count >= 400:
                    page_str = self.clean_text("\n\n".join(current_page_text))
                    pages.append(DocumentPage(page_number=page_number, text=page_str))
                    current_page_text = []
                    current_word_count = 0
                    page_number += 1

            if current_page_text:
                page_str = self.clean_text("\n\n".join(current_page_text))
                pages.append(DocumentPage(page_number=page_number, text=page_str))

            total_pages = max(len(pages), 1)
            return DocumentContent(
                title=title,
                pages=pages,
                total_pages=total_pages,
                metadata=metadata or {}
            )
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX '{title}': {e}", exc_info=True)
            raise ValueError(f"Error parsing DOCX document: {str(e)}")

    def extract_document(
        self,
        file_bytes: bytes,
        file_name: str,
        title: str = "",
        metadata: Optional[Dict[str, Any]] = None
    ) -> DocumentContent:
        """Dispatches extraction based on file extension and MIME heuristics."""
        lower_name = file_name.lower()
        if lower_name.endswith(".pdf"):
            return self.extract_from_pdf(file_bytes, title=title or file_name, metadata=metadata)
        elif lower_name.endswith((".docx", ".doc")):
            return self.extract_from_docx(file_bytes, title=title or file_name, metadata=metadata)
        else:
            raise ValueError(f"Unsupported document format for '{file_name}'. Only PDF and DOCX are supported.")
